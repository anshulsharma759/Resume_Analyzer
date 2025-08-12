import streamlit as st
import pdfplumber
import docx
import os
import re
import requests
import io
import spacy
from spacy.matcher import PhraseMatcher

# Helper function to extract text from PDF
def extract_text_from_pdf(file):
    try:
        with pdfplumber.open(file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if not text.strip():
            return "No extractable text found. If this is a scanned PDF, try OCR or upload a DOCX."
        return text
    except Exception:
        return "Error reading PDF. Try converting to DOCX."

# Helper function to extract text from DOCX
def extract_text_from_docx(file):
    # Accept both file-like and path
    if hasattr(file, 'read'):
        file.seek(0)
        doc = docx.Document(io.BytesIO(file.read()))
    else:
        doc = docx.Document(file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join([t for t in text if t.strip()])

# Technical skills and related roles/suggestions
TECH_SKILLS = {
    "python": {"roles": ["Data Scientist", "Backend Developer", "ML Engineer"], "upskill": ["Deep Learning", "FastAPI", "Data Engineering"]},
    "java": {"roles": ["Backend Developer", "Android Developer"], "upskill": ["Spring Boot", "Microservices"]},
    "c++": {"roles": ["Embedded Engineer", "Game Developer"], "upskill": ["Modern C++", "Concurrency"]},
    "sql": {"roles": ["Data Analyst", "Database Admin"], "upskill": ["NoSQL", "Data Warehousing"]},
    "excel": {"roles": ["Business Analyst"], "upskill": ["Power BI", "VBA"]},
    "javascript": {"roles": ["Frontend Developer", "Full Stack Developer"], "upskill": ["React", "Node.js"]},
    "aws": {"roles": ["Cloud Engineer", "DevOps Engineer"], "upskill": ["Serverless", "Kubernetes"]},
    "docker": {"roles": ["DevOps Engineer"], "upskill": ["Kubernetes", "CI/CD"]},
    "react": {"roles": ["Frontend Developer"], "upskill": ["Next.js", "TypeScript"]},
    "machine learning": {"roles": ["ML Engineer", "Data Scientist"], "upskill": ["Deep Learning", "MLOps"]},
    "typescript": {"roles": ["Frontend Developer", "Full Stack Developer"], "upskill": ["React", "Node.js"]},
    "html": {"roles": ["Frontend Developer"], "upskill": ["CSS", "JavaScript"]},
    "css": {"roles": ["Frontend Developer"], "upskill": ["Sass", "Tailwind"]},
    "node.js": {"roles": ["Backend Developer", "Full Stack Developer"], "upskill": ["Express.js", "TypeScript"]},
    "mongodb": {"roles": ["Database Admin", "Backend Developer"], "upskill": ["Sharding", "Replication"]},
    "git": {"roles": ["DevOps Engineer", "Developer"], "upskill": ["GitHub Actions", "CI/CD"]},
    "linux": {"roles": ["DevOps Engineer", "SysAdmin"], "upskill": ["Shell Scripting", "Security"]},
    "power bi": {"roles": ["Data Analyst"], "upskill": ["DAX", "Power Query"]},
    "tableau": {"roles": ["Data Analyst"], "upskill": ["Dashboard Design", "Data Prep"]},
    "azure": {"roles": ["Cloud Engineer", "DevOps Engineer"], "upskill": ["Azure DevOps", "Functions"]},
    "kubernetes": {"roles": ["DevOps Engineer", "Cloud Engineer"], "upskill": ["Helm", "Service Mesh"]},
    # ...add more as needed
}

# Proficiency keywords
PROFICIENCY_LEVELS = [
    ("expert", 3),
    ("advanced", 3),
    ("proficient", 2),
    ("intermediate", 2),
    ("familiar", 1),
    ("beginner", 1),
    ("basic", 1),
]

def detect_proficiency(skill, text):
    for word, level in PROFICIENCY_LEVELS:
        pattern = rf"\\b{word}\\b.*?\\b{skill}\\b|\\b{skill}\\b.*?\\b{word}\\b"
        if re.search(pattern, text, re.IGNORECASE):
            return word.capitalize(), level
    # Default: found but no level
    return "Mentioned", 1

def analyze_technical_skills(text):
    found = []
    for skill in TECH_SKILLS:
        if re.search(rf"\\b{re.escape(skill)}\\b", text, re.IGNORECASE):
            prof, level = detect_proficiency(skill, text)
            found.append((skill, prof, level))
    return found

def suggest_next_role_and_upskill(skills):
    roles = set()
    upskills = set()
    for skill, _, _ in skills:
        info = TECH_SKILLS.get(skill.lower())
        if info:
            roles.update(info["roles"])
            upskills.update(info["upskill"])
    return list(roles), list(upskills)

def extract_possible_skills(text):
    # Extract capitalized words, words with special chars, and common tech patterns
    # This is a heuristic and will catch most tech skills
    words = re.findall(r"\\b([A-Z][a-zA-Z0-9\+\-\.\#]*)\\b", text)
    # Add lowercase techs (python, sql, etc.)
    common_techs = re.findall(r"\\b([a-z]{2,}\+?\#?)\\b", text)
    # Merge and deduplicate
    all_skills = set(words + common_techs)
    # Remove common English words (stopwords)
    stopwords = set(["and", "the", "with", "for", "from", "that", "this", "have", "has", "are", "was", "will", "can", "not", "but", "you", "your", "our", "their", "they", "she", "him", "her", "his", "its", "who", "what", "when", "where", "how", "which", "also", "use", "using", "used", "work", "worked", "working", "project", "projects", "team", "teams", "company", "companies", "experience", "years", "month", "months", "year", "role", "roles", "responsible", "responsibility", "responsibilities", "etc"])
    filtered_skills = [w for w in all_skills if w.lower() not in stopwords and len(w) > 1]
    return sorted(filtered_skills, key=lambda w: text.lower().count(w.lower()), reverse=True)

def get_skills_from_api(keywords):
    """
    Uses the DataAtWork Skills API to extract skills from a list of keywords.
    API: https://skills.emsidata.com/ (free, no key required)
    """
    url = "https://skills.emsidata.com/skills/auto-complete"
    found_skills = set()
    for word in keywords:
        params = {"q": word, "limit": 1}
        try:
            resp = requests.get(url, params=params, timeout=2)
            if resp.status_code == 200:
                data = resp.json()
                if data and "data" in data and data["data"]:
                    found_skills.add(data["data"][0]["name"])
        except Exception:
            continue
    return sorted(found_skills)

def estimate_proficiency(skill, text):
    # Look for context clues for proficiency
    prof_map = {
        'expert': 3, 'advanced': 3, 'proficient': 2, 'intermediate': 2, 'familiar': 1, 'beginner': 1, 'basic': 1
    }
    for prof, level in prof_map.items():
        if re.search(rf"\\b{prof}\\b.*?\\b{skill}\\b|\\b{skill}\\b.*?\\b{prof}\\b", text, re.IGNORECASE):
            return prof.capitalize(), level
    # Look for years of experience
    match = re.search(rf"(\\d+)\\+?\\s*(years|yrs|year|yr)[^\\.]*?\\b{skill}\\b", text, re.IGNORECASE)
    if match:
        years = int(match.group(1))
        if years >= 5:
            return f"{years} years (Expert)", 3
        elif years >= 2:
            return f"{years} years (Intermediate)", 2
        else:
            return f"{years} years (Beginner)", 1
    return "Mentioned", 1

def get_roles_and_upskills(skills):
    # Use a static mapping for demo, can be replaced with API
    role_map = {
        'python': ['Data Scientist', 'Backend Developer', 'ML Engineer'],
        'java': ['Backend Developer', 'Android Developer'],
        'sql': ['Data Analyst', 'Database Admin'],
        'javascript': ['Frontend Developer', 'Full Stack Developer'],
        'aws': ['Cloud Engineer', 'DevOps Engineer'],
        'docker': ['DevOps Engineer'],
        'react': ['Frontend Developer'],
        'machine learning': ['ML Engineer', 'Data Scientist'],
        # ...add more as needed
    }
    upskill_map = {
        'python': ['Deep Learning', 'FastAPI', 'Data Engineering'],
        'java': ['Spring Boot', 'Microservices'],
        'sql': ['NoSQL', 'Data Warehousing'],
        'javascript': ['React', 'Node.js'],
        'aws': ['Serverless', 'Kubernetes'],
        'docker': ['Kubernetes', 'CI/CD'],
        'react': ['Next.js', 'TypeScript'],
        'machine learning': ['Deep Learning', 'MLOps'],
        # ...add more as needed
    }
    roles, upskills = set(), set()
    for skill in skills:
        s = skill.lower()
        if s in role_map:
            roles.update(role_map[s])
        if s in upskill_map:
            upskills.update(upskill_map[s])
    return list(roles), list(upskills)

def get_resume_tips(skills, upskills):
    tips = []
    if not skills:
        tips.append("Add more technical skills relevant to your target job.")
    if upskills:
        tips.append(f"Consider upskilling in: {', '.join(upskills)}.")
    tips.append("Highlight your proficiency and years of experience for each skill.")
    tips.append("Tailor your resume for the job role you want.")
    return tips

def extract_skill_candidates(text):
    # Extract capitalized words, tech patterns, and multi-word phrases
    cap_words = re.findall(r"\b([A-Z][a-zA-Z0-9\+\-\.\#]{1,})\b", text)
    tech_words = re.findall(r"\b([a-z]{2,}\+?\#?)\b", text)
    phrases = re.findall(r"([A-Za-z]{2,}(?: [A-Za-z]{2,}){1,2})", text)
    all_candidates = set(cap_words + tech_words + phrases)
    stopwords = set(["and", "the", "with", "for", "from", "that", "this", "have", "has", "are", "was", "will", "can", "not", "but", "you", "your", "our", "their", "they", "she", "him", "her", "his", "its", "who", "what", "when", "where", "how", "which", "also", "use", "using", "used", "work", "worked", "working", "project", "projects", "team", "teams", "company", "companies", "experience", "years", "month", "months", "year", "role", "roles", "responsible", "responsibility", "responsibilities", "etc"])
    filtered = [w for w in all_candidates if w.lower() not in stopwords and len(w) > 1 and not w.isdigit()]
    return sorted(filtered, key=lambda w: text.lower().count(w.lower()), reverse=True)

def extract_technical_words(text):
    # List of common technical keywords and patterns (expand as needed)
    tech_patterns = [
        r"crm", r"cloud", r"database", r"sql", r"python", r"java", r"c\+\+", r"c#", r"javascript", r"react", r"node", r"aws", r"azure", r"gcp", r"docker", r"kubernetes", r"devops", r"etl", r"data", r"analytics", r"machine learning", r"ai", r"ml", r"nlp", r"power bi", r"tableau", r"snowflake", r"hadoop", r"spark", r"sas", r"siebel", r"salesforce", r"sap", r"oracle", r"linux", r"unix", r"windows", r"git", r"jira", r"agile", r"scrum", r"ci/cd", r"microservices", r"rest", r"graphql", r"api", r"big data", r"nosql", r"mongodb", r"postgres", r"mysql", r"db2", r"ssis", r"ssrs", r"ssas", r"dax", r"matlab", r"r studio", r"scala", r"go", r"typescript", r"html", r"css", r"bootstrap", r"django", r"flask", r"fastapi", r"dotnet", r".net", r"vb.net", r"vbscript", r"powershell", r"bash", r"shell", r"unix", r"linux", r"windows", r"vmware", r"hyper-v", r"ansible", r"puppet", r"chef", r"jenkins", r"terraform", r"snowflake", r"siebel crm"
    ]
    # Lowercase text for matching
    text_lc = text.lower()
    found = set()
    for pat in tech_patterns:
        if re.search(rf"\\b{pat}\\b", text_lc):
            found.add(pat)
    # Also add all unique words (except stopwords) for API
    words = set(re.findall(r"[a-zA-Z0-9\+\-\.\#]{2,}", text_lc))
    stopwords = set(["and", "the", "with", "for", "from", "that", "this", "have", "has", "are", "was", "will", "can", "not", "but", "you", "your", "our", "their", "they", "she", "him", "her", "his", "its", "who", "what", "when", "where", "how", "which", "also", "use", "using", "used", "work", "worked", "working", "project", "projects", "team", "teams", "company", "companies", "experience", "years", "month", "months", "year", "role", "roles", "responsible", "responsibility", "responsibilities", "etc"])
    tech_words = [w for w in words if w not in stopwords]
    # Merge and deduplicate
    return sorted(found.union(tech_words), key=lambda w: text_lc.count(w), reverse=True)

def extract_technical_keywords(text):
    import itertools
    text_lc = text.lower()
    # Tokenize and keep original for capitalization
    tokens = re.findall(r"[a-zA-Z0-9\+\-\.\#]{2,}", text)
    tokens_lc = [t.lower() for t in tokens]
    stopwords = set(["and", "the", "with", "for", "from", "that", "this", "have", "has", "are", "was", "will", "can", "not", "but", "you", "your", "our", "their", "they", "she", "him", "her", "his", "its", "who", "what", "when", "where", "how", "which", "also", "use", "using", "used", "work", "worked", "working", "project", "projects", "team", "teams", "company", "companies", "experience", "years", "month", "months", "year", "role", "roles", "responsible", "responsibility", "responsibilities", "etc"])
    # Remove stopwords and short tokens
    tokens = [t for t in tokens if t.lower() not in stopwords and len(t) > 2 and not t.isdigit()]
    # Build n-grams (1-4 words)
    ngrams = set()
    for n in range(1, 5):
        for i in range(len(tokens) - n + 1):
            phrase = " ".join(tokens[i:i+n])
            # Heuristic: likely technical if contains numbers, +, #, or is long enough
            if (any(c in phrase for c in "+#") or len(phrase) > 4 or n > 1):
                ngrams.add(phrase)
    # Tech patterns and dictionary
    tech_patterns = [
        "crm", "cloud", "database", "sql", "python", "java", "c++", "c#", "javascript", "react", "node", "aws", "azure", "gcp", "docker", "kubernetes", "devops", "etl", "data", "analytics", "machine learning", "ai", "ml", "nlp", "power bi", "tableau", "snowflake", "hadoop", "spark", "sas", "siebel", "salesforce", "sap", "oracle", "linux", "unix", "windows", "git", "jira", "agile", "scrum", "ci/cd", "microservices", "rest", "graphql", "api", "big data", "nosql", "mongodb", "postgres", "mysql", "db2", "ssis", "ssrs", "ssas", "dax", "matlab", "r studio", "scala", "go", "typescript", "html", "css", "bootstrap", "django", "flask", "fastapi", "dotnet", ".net", "vb.net", "vbscript", "powershell", "bash", "shell", "vmware", "hyper-v", "ansible", "puppet", "chef", "jenkins", "terraform", "snowflake", "siebel crm"
    ]
    # Use a set of known tech words for direct matching
    tech_dict = set(tech_patterns)
    good_keywords = set()
    for phrase in ngrams:
        for pat in tech_patterns:
            if pat in phrase:
                good_keywords.add(phrase)
    # Add ngrams that are 2-4 words and not in stopwords
    for phrase in ngrams:
        if len(phrase.split()) > 1 and phrase not in good_keywords:
            good_keywords.add(phrase)
    # Add single tokens that are in the tech dictionary
    for t in tokens:
        if t.lower() in tech_dict:
            good_keywords.add(t)
    # Add capitalized words (likely proper nouns/tech)
    for t in tokens:
        if t[0].isupper() and t.lower() not in stopwords:
            good_keywords.add(t)
    # Remove duplicates and sort by frequency in text
    return sorted(good_keywords, key=lambda w: text_lc.count(w.lower()), reverse=True)

def extract_skills_spacy(text):
    base_skills = [
        "python", "java", "c++", "c#", "sql", "javascript", "react", "node.js", "aws", "azure", "gcp", "docker", "kubernetes", "devops", "etl", "data analytics", "machine learning", "ai", "ml", "nlp", "power bi", "tableau", "snowflake", "hadoop", "spark", "sas", "siebel", "siebel crm", "salesforce", "sap", "oracle", "linux", "unix", "windows", "git", "jira", "agile", "scrum", "ci/cd", "microservices", "rest api", "graphql", "api", "big data", "nosql", "mongodb", "postgres", "mysql", "db2", "ssis", "ssrs", "ssas", "dax", "matlab", "r studio", "scala", "go", "typescript", "html", "css", "bootstrap", "django", "flask", "fastapi", ".net", "vb.net", "vbscript", "powershell", "bash", "shell", "vmware", "hyper-v", "ansible", "puppet", "chef", "jenkins", "terraform"
    ]
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    noun_phrases = set(chunk.text.strip() for chunk in doc.noun_chunks if len(chunk.text.strip()) > 2)
    cap_words = set([token.text for token in doc if token.is_title and len(token.text) > 2])
    # Remove irrelevant patterns (emails, locations, mobile, personal info, etc.)
    irrelevant_patterns = [
        r"[\w\.-]+@[\w\.-]+",  # email
        r"\b(?:city|state|country|india|indian|usa|canada|uk|london|new york|bangalore|delhi|mumbai|chennai|hyderabad|pune|california|texas|paris|berlin|tokyo|singapore|dubai|sydney|melbourne|toronto|vancouver|noida)\b",  # locations
        r"\b(?:nationality|management summary|mobile|residence|birth|training & certification|training|certification|summary|profile|personal|address|contact|phone|email|dob|date of birth|gender|marital|father|mother|passport|pan|aadhaar|religion|caste|category|languages|language|hobbies|interests|objective|career objective|career summary|about|bio|curriculum vitae|cv|resume|details|declaration|place|date|signature|references?)\b",  # personal info
        r"\b\d{10,}\b",  # mobile numbers
    ]
    def is_irrelevant(s):
        for pat in irrelevant_patterns:
            if re.search(pat, s, re.IGNORECASE):
                return True
        return False
    all_skills = set(base_skills) | noun_phrases | cap_words
    stopwords = set(["and", "the", "with", "for", "from", "that", "this", "have", "has", "are", "was", "will", "can", "not", "but", "you", "your", "our", "their", "they", "she", "him", "her", "his", "its", "who", "what", "when", "where", "how", "which", "also", "use", "using", "used", "work", "worked", "working", "project", "projects", "team", "teams", "company", "companies", "experience", "years", "month", "months", "year", "role", "roles", "responsible", "responsibility", "responsibilities", "etc"])
    all_skills = [s for s in all_skills if s.lower() not in stopwords and len(s) > 2 and not is_irrelevant(s)]
    matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
    patterns = [nlp.make_doc(skill) for skill in all_skills]
    matcher.add("SKILLS", patterns)
    matches = matcher(doc)
    found_skills = set([doc[start:end].text for match_id, start, end in matches])
    return sorted(found_skills, key=lambda w: text.lower().count(w.lower()), reverse=True)

def suggest_roles_from_skills(skills):
    # Simple mapping for demonstration
    role_map = {
        "python": ["Data Scientist", "Backend Developer", "ML Engineer"],
        "java": ["Backend Developer", "Android Developer"],
        "sql": ["Data Analyst", "Database Admin"],
        "javascript": ["Frontend Developer", "Full Stack Developer"],
        "aws": ["Cloud Engineer", "DevOps Engineer"],
        "docker": ["DevOps Engineer"],
        "react": ["Frontend Developer"],
        "machine learning": ["ML Engineer", "Data Scientist"],
        "snowflake": ["Data Engineer", "Cloud Data Engineer"],
        "siebel crm": ["CRM Consultant", "Siebel Developer"],
        "power bi": ["BI Developer", "Data Analyst"],
        "kubernetes": ["DevOps Engineer", "Cloud Engineer"],
        "salesforce": ["Salesforce Developer", "CRM Consultant"],
        # ...add more as needed
    }
    roles = set()
    for skill in skills:
        for k, v in role_map.items():
            if k in skill.lower():
                roles.update(v)
    return sorted(roles)

def suggest_upskills_from_cluster(skills):
    # Map clusters to upskill suggestions
    cluster_map = {
        'data': ["python", "sql", "power bi", "tableau", "hadoop", "spark", "machine learning", "snowflake", "data engineering"],
        'cloud': ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "devops"],
        'web': ["javascript", "react", "node.js", "typescript", "html", "css", "django", "flask", "fastapi"],
        'crm': ["salesforce", "siebel crm", "sap", "oracle crm"],
        'bi': ["power bi", "tableau", "ssis", "ssrs", "ssas", "dax"],
        'devops': ["docker", "kubernetes", "jenkins", "ansible", "puppet", "terraform", "ci/cd"],
        # ...add more clusters as needed
    }
    # Detect clusters present in user's skills
    cluster_hits = {k: set(v) & set(map(str.lower, skills)) for k, v in cluster_map.items()}
    # Find the cluster with the most overlap
    best_cluster = max(cluster_hits.items(), key=lambda x: len(x[1]), default=(None, set()))[0]
    if not best_cluster:
        return [], None
    # Suggest skills from the cluster not already present
    suggested = [s for s in cluster_map[best_cluster] if s.lower() not in map(str.lower, skills)]
    return suggested, best_cluster

st.title("Resume Analyzer: spaCy Skill Extraction & Role Suggestion")
st.write("Upload your resume (PDF or DOCX) to extract technical skills and get suggested tech roles.")

uploaded_file = st.file_uploader("Choose a resume file", type=["pdf", "docx"])

if uploaded_file:
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file type.")
        text = ""

    if text:
        st.subheader("Technical Skills Detected by spaCy:")
        skills = extract_skills_spacy(text)
        st.write(", ".join(skills) if skills else "No technical skills detected.")
        st.subheader("Suggested Tech Roles:")
        roles = suggest_roles_from_skills(skills)
        st.write(", ".join(roles) if roles else "No role suggestions found. Add more technical skills to your resume.")
        # New section: Upskill suggestions based on tech cluster
        upskills, cluster = suggest_upskills_from_cluster(skills)
        st.subheader("What Skills to Add Next (Based on Your Tech Cluster):")
        if upskills:
            st.info(f"Based on your main tech cluster ({cluster.title()}), consider adding: {', '.join(upskills)}")
        else:
            st.info("No upskill suggestions found. Try adding more core technical skills to your resume.")
    else:
        st.error("Could not extract text from the file.")
